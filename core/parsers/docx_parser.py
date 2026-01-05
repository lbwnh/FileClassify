"""
DOCX 解析器实现。

本模块提供 DOCX/DOC 文件解析功能。
"""

from .base import BaseParser
from typing import Dict


class DOCXParser(BaseParser):
    """
    DOCX 和 DOC 文件解析器。

    使用 python-docx 进行文本提取。
    """
    
    def extract_text(self) -> str:
        """
        从 DOCX 文件中提取文本内容。

        返回:
            str: 提取的文本内容
        """
        try:
            import docx
            
            doc = docx.Document(self.file_path)
            
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return '\n'.join(text_content)
        
        except ImportError:
            return f"[DOCX Parser] Cannot extract text: python-docx not installed.\nFile: {self.file_path.name}"
        
        except Exception as e:
            return f"[DOCX Parser] Error extracting text: {str(e)}\nFile: {self.file_path.name}"
    
    def extract_metadata(self) -> Dict[str, any]:
        """
        从 DOCX 文件中提取元数据。

        返回:
            dict: DOCX 元数据
        """
        metadata = {
            'title': None,
            'author': None,
            'subject': None,
            'keywords': None,
            'created_date': None,
            'modified_date': None,
            'last_modified_by': None,
            'paragraph_count': 0,
            'word_count': 0
        }
        
        try:
            import docx
            
            doc = docx.Document(self.file_path)
            
            core_props = doc.core_properties
            
            metadata['title'] = core_props.title
            metadata['author'] = core_props.author
            metadata['subject'] = core_props.subject
            metadata['keywords'] = core_props.keywords
            metadata['created_date'] = core_props.created
            metadata['modified_date'] = core_props.modified
            metadata['last_modified_by'] = core_props.last_modified_by
            
            metadata['paragraph_count'] = len(doc.paragraphs)
            
            text = self.extract_text()
            metadata['word_count'] = len(text.split())
        
        except Exception as e:
            metadata['error'] = str(e)
        
        return metadata
